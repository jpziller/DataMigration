"""Solution document generator (roadmap #4).

Auto-drafts a migration solution/design document from data this framework
already has -- load-order analysis (load_order.py), a filled-in mapping doc
(mapping_doc.py), and profiling data (profiling.py) -- instead of writing one
by hand for every project.

The default document is built entirely in Python (see _build_default_docx)
-- there is no binary template checked into git. This follows the same
principle as auto_mapper.py's thesaurus: the framework's own baseline
content is reviewable, diffable text, not an opaque file. It covers, in
order: what's being built (object scope), how it's being done (the SQL-
centric methodology, in plain language), the load order, one section per
object (source table, row count, field-mapping status, profiling summary),
and an optional appendix with the full field-by-field mapping detail.

A data architect at a different company can still bring their own branding
(logo, colors, house style) by supplying --template with a .docx built in
Word that contains the same context fields as Jinja2 tags (via `docxtpl`).
If no --template is given, the framework falls back to the default. The
context dict handed to either path is identical -- see gather_context()'s
return shape below. A custom template needs these tags to render fully:

    {{ project_name }}, {{ company_name }}, {{ prepared_by }},
    {{ prepared_date }}, {{ target_org_alias }}

    {% for o in objects %}
      {{ o.name }} {{ o.load_level }} {{ o.load_sequence }}
      {{ o.source_table }} {{ o.row_count }}
      {{ o.mapping.total_source_fields }} {{ o.mapping.mapped }}
      {{ o.mapping.migrate_yes }} {{ o.mapping.migrate_no }}
      {{ o.mapping.migrate_review }} {{ o.mapping.migrate_blank }}
      {{ o.profile_summary.fields_profiled }}
      {{ o.profile_summary.avg_populated_pct }}
      {{ o.profile_summary.low_population_count }}
    {% endfor %}

    {% if unresolved_cycles %} ... {% endif %}
    {% if include_appendix %}
      {% for r in appendix_rows %}
        {{ r.object }} {{ r.source_field }} {{ r.target_field }}
        {{ r.migrate }} {{ r.notes }}
      {% endfor %}
    {% endif %}

`o.mapping` and `o.profile_summary` are None for an object with no mapping
sheet / no profiling data yet -- guard with `{% if o.mapping %}` in a custom
template. Word autocorrects straight quotes to curly quotes as you type,
which breaks Jinja tags typed directly into a document -- disable
autocorrect (or paste tags in) before saving a custom template.
"""
import os
import re
from datetime import datetime

import openpyxl
from docx import Document
from docxtpl import DocxTemplate
from sqlalchemy import text

import load_order as lo

_INVALID_SHEET_CHARS = re.compile(r"[:\\/?*\[\]]")

# Mirrors mapping_doc.py's column layout -- see that module's _HEADERS.
_COL_SOURCE_FIELD_API = 2
_COL_SOURCE_NOTES = 8
_COL_MIGRATE_DATA = 9
_COL_TARGET_FIELD_API = 15

_LOW_POPULATION_PCT = 20.0


def _safe_sheet_name(name):
    return _INVALID_SHEET_CHARS.sub("_", name)[:31]


def _table_exists(cx, schema, table):
    return cx.execute(text("SELECT OBJECT_ID(:t, 'U')"), {"t": f"{schema}.{table}"}).scalar() is not None


def _mapping_summary_for_sheet(ws, include_appendix, object_name, appendix_rows):
    source_table = ws.cell(row=1, column=2).value
    total = mapped = yes = no = review = blank = 0

    for row in ws.iter_rows(min_row=4):
        source_field = row[_COL_SOURCE_FIELD_API - 1].value
        if not source_field:
            continue
        total += 1

        target_field = row[_COL_TARGET_FIELD_API - 1].value
        if target_field and str(target_field).strip():
            mapped += 1

        migrate = row[_COL_MIGRATE_DATA - 1].value
        migrate = str(migrate).strip() if migrate else ""
        if migrate == "Yes":
            yes += 1
        elif migrate == "No":
            no += 1
        elif migrate == "Review":
            review += 1
        else:
            blank += 1

        if include_appendix:
            appendix_rows.append({
                "object": object_name,
                "source_field": source_field,
                "target_field": target_field or "",
                "migrate": migrate,
                "notes": row[_COL_SOURCE_NOTES - 1].value or "",
            })

    summary = {
        "total_source_fields": total,
        "mapped": mapped,
        "migrate_yes": yes,
        "migrate_no": no,
        "migrate_review": review,
        "migrate_blank": blank,
    }
    return source_table, summary


def _profile_summary_for_table(cx, schema, table_name):
    if not table_name or not _table_exists(cx, schema, "FieldProfile"):
        return None
    pcts = [
        r[0] for r in cx.execute(
            text(f"SELECT PopulatedPct FROM [{schema}].[FieldProfile] WHERE ObjectOrTable = :t"),
            {"t": table_name},
        ).fetchall()
        if r[0] is not None
    ]
    if not pcts:
        return None
    return {
        "fields_profiled": len(pcts),
        "avg_populated_pct": sum(pcts) / len(pcts),
        "low_population_count": sum(1 for p in pcts if p < _LOW_POPULATION_PCT),
    }


def gather_context(sf, engine, object_names, mapping_path=None, schema="dbo",
                    company_name=None, project_name=None, prepared_by=None,
                    target_org_alias=None, include_appendix=False):
    """Assemble the context dict rendered into either the default document
    or a custom docxtpl template -- see this module's docstring for the
    full tag contract."""
    object_names = list(object_names)
    load_result = lo.analyze_load_order(sf, engine, object_names, schema=schema)
    order_by_object = {row["object"]: row for row in load_result["order"]}
    self_ref = load_result["self_references"]

    cycle_by_object = {}
    for group in load_result["unresolved_cycles"]:
        for name in group:
            cycle_by_object[name] = [m for m in group if m != name]

    mapping_sheets = {}
    if mapping_path and os.path.exists(mapping_path):
        wb = openpyxl.load_workbook(mapping_path, data_only=True)
        for name in object_names:
            sheet_name = _safe_sheet_name(name)
            if sheet_name in wb.sheetnames:
                mapping_sheets[name] = wb[sheet_name]

    appendix_rows = []
    objects = []

    with engine.connect() as cx:
        for name in object_names:
            order_row = order_by_object.get(name)

            source_table = None
            mapping_summary = None
            ws = mapping_sheets.get(name)
            if ws is not None:
                source_table, mapping_summary = _mapping_summary_for_sheet(
                    ws, include_appendix, name, appendix_rows
                )

            row_count = None
            if source_table and _table_exists(cx, schema, source_table):
                row_count = cx.execute(text(f"SELECT COUNT(*) FROM [{schema}].[{source_table}]")).scalar()

            profile_summary = _profile_summary_for_table(cx, schema, source_table or name)

            objects.append({
                "name": name,
                "load_level": order_row["level"] if order_row else None,
                "load_sequence": order_row["sequence"] if order_row else None,
                "self_reference_fields": self_ref.get(name),
                "in_cycle": name in cycle_by_object,
                "cycle_members": cycle_by_object.get(name),
                "source_table": source_table,
                "row_count": row_count,
                "mapping": mapping_summary,
                "profile_summary": profile_summary,
            })

    objects.sort(key=lambda o: (o["load_sequence"] is None, o["load_sequence"] or 0))

    return {
        "company_name": company_name or "",
        "project_name": project_name or "Salesforce Data Migration",
        "prepared_by": prepared_by or "",
        "prepared_date": datetime.now().strftime("%Y-%m-%d"),
        "target_org_alias": target_org_alias or "",
        "objects": objects,
        "unresolved_cycles": load_result["unresolved_cycles"],
        "include_appendix": include_appendix,
        "appendix_rows": appendix_rows,
    }


def _add_bold(paragraph, text_value):
    run = paragraph.add_run(text_value)
    run.bold = True
    return run


def _object_notes(obj):
    notes = []
    if obj["self_reference_fields"]:
        notes.append(f"self-references via {', '.join(obj['self_reference_fields'])} (two-pass load)")
    if obj["in_cycle"]:
        notes.append(f"circular dependency with {', '.join(obj['cycle_members'])}")
    return "; ".join(notes) if notes else ""


def _build_default_docx(context, output_path):
    doc = Document()

    doc.add_heading(context["project_name"], level=0)
    if context["company_name"]:
        p = doc.add_paragraph()
        _add_bold(p, context["company_name"])
    meta_bits = []
    if context["prepared_by"]:
        meta_bits.append(f"Prepared by: {context['prepared_by']}")
    meta_bits.append(f"Date: {context['prepared_date']}")
    if context["target_org_alias"]:
        meta_bits.append(f"Target org: {context['target_org_alias']}")
    doc.add_paragraph(" | ".join(meta_bits))
    doc.add_page_break()

    # 1. What we are building
    doc.add_heading("1. What We Are Building", level=1)
    doc.add_paragraph(
        f"{context['project_name']} migrates data into "
        f"{context['target_org_alias'] or 'the target Salesforce org'} for the "
        "following object(s), listed in the order they will be loaded so that "
        "every parent record exists before its children are inserted:"
    )
    for obj in context["objects"]:
        level = f" (level {obj['load_level']})" if obj["load_level"] is not None else ""
        doc.add_paragraph(f"{obj['name']}{level}", style="List Bullet")
    if context["unresolved_cycles"]:
        doc.add_paragraph(
            "The following objects have a circular dependency that could not be "
            "resolved automatically and will need a two-pass load or manual "
            "sequencing:"
        )
        for group in context["unresolved_cycles"]:
            doc.add_paragraph(", ".join(group), style="List Bullet")

    # 2. How we are doing this
    doc.add_heading("2. How We Are Doing This", level=1)
    doc.add_paragraph(
        "This migration follows a SQL-centric approach: SQL Server acts as the "
        "integration hub between the source data and the target Salesforce org, "
        "rather than moving data directly between the two systems. Every object "
        "passes through three stages -- replicate, transform, and load -- each "
        "independently reviewable before anything touches production data."
    )
    doc.add_paragraph(
        "Replicate: source data is pulled into a mirror database in SQL Server, "
        "preserving the original values so later steps always have an "
        "unmodified reference to fall back on."
    )
    doc.add_paragraph(
        "Transform: the logic that decides how each source field becomes a "
        "target field -- cleansing, splitting, lookups, defaulting -- is "
        "written as versioned SQL, one script per object, producing a load "
        "table that mirrors exactly what will be submitted to Salesforce."
    )
    doc.add_paragraph(
        "Load: the load table is submitted to Salesforce in bulk. Because the "
        "API returns successful and failed records as separate, unordered "
        "sets, each submitted row is matched back to its result by "
        "fingerprinting its business columns rather than assuming row order -- "
        "inserts carry a dedicated migration-key field for this purpose. Every "
        "load table is also sorted so that records sharing the same parent "
        "land in the same processing batch, and checked for duplicate or "
        "missing migration keys, before it is ever submitted."
    )

    # 3. Migration scope & load order
    doc.add_heading("3. Migration Scope & Load Order", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "Object", "Load Level", "Load Sequence", "Notes"
    for obj in context["objects"]:
        cells = table.add_row().cells
        cells[0].text = obj["name"]
        cells[1].text = "" if obj["load_level"] is None else str(obj["load_level"])
        cells[2].text = "" if obj["load_sequence"] is None else str(obj["load_sequence"])
        cells[3].text = _object_notes(obj)

    # 4. Object details
    doc.add_heading("4. Object Details", level=1)
    for obj in context["objects"]:
        doc.add_heading(obj["name"], level=2)
        if obj["source_table"]:
            count_text = f" ({obj['row_count']:,} row(s))" if obj["row_count"] is not None else ""
            doc.add_paragraph(f"Source table: {obj['source_table']}{count_text}")

        m = obj["mapping"]
        if m:
            doc.add_paragraph(
                f"Field mapping: {m['mapped']} of {m['total_source_fields']} source field(s) "
                f"mapped to a target field. {m['migrate_yes']} recommended to migrate, "
                f"{m['migrate_no']} not recommended, {m['migrate_review']} flagged for review, "
                f"{m['migrate_blank']} not yet decided."
            )
        else:
            doc.add_paragraph(
                "Field mapping has not yet been documented for this object -- run "
                "generate-mapping-doc and auto-map to populate it."
            )

        p = obj["profile_summary"]
        if p:
            doc.add_paragraph(
                f"Data profiling: {p['fields_profiled']} field(s) profiled, "
                f"{p['avg_populated_pct']:.1f}% average population, "
                f"{p['low_population_count']} field(s) below {_LOW_POPULATION_PCT:.0f}% populated."
            )
        else:
            doc.add_paragraph("Data profiling: not yet run for this object's source table.")

    # 5. Appendix (only when explicitly requested)
    if context["include_appendix"]:
        doc.add_heading("5. Appendix: Full Field Mapping", level=1)
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text, hdr[4].text = (
            "Object", "Source Field", "Target Field", "Migrate", "Notes"
        )
        for r in context["appendix_rows"]:
            cells = table.add_row().cells
            cells[0].text = r["object"]
            cells[1].text = r["source_field"]
            cells[2].text = r["target_field"]
            cells[3].text = r["migrate"]
            cells[4].text = r["notes"]

    doc.save(output_path)


def generate_solution_doc(sf, engine, output_path, object_names, mapping_path=None,
                           template_path=None, schema="dbo", company_name=None,
                           project_name=None, prepared_by=None, target_org_alias=None,
                           include_appendix=False):
    context = gather_context(
        sf, engine, object_names, mapping_path=mapping_path, schema=schema,
        company_name=company_name, project_name=project_name, prepared_by=prepared_by,
        target_org_alias=target_org_alias, include_appendix=include_appendix,
    )

    if template_path:
        tpl = DocxTemplate(template_path)
        tpl.render(context)
        tpl.save(output_path)
    else:
        _build_default_docx(context, output_path)

    return output_path, context
